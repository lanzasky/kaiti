from __future__ import annotations

import argparse
import csv
import json
import os
import re
import urllib.request
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise RuntimeError(
        "kaiti_skill requires python-docx. Install it with `python -m pip install python-docx`."
    ) from exc


DEFAULT_TEMPLATE_STRUCTURE = [
    "第一部分 立论依据",
    "1.1 研究背景与意义",
    "1.2 国内外研究现状",
    "1.3 文献评述与项目切入点",
    "主要参考文献",
    "第二部分 研究方案",
    "2.1 研究目标",
    "2.2 研究内容",
    "2.3 研究技术路线",
    "2.4 实验方案",
    "2.5 可行性分析",
    "2.6 预期研究进展",
    "三、研究基础",
    "四、经费预算",
]


SYNTHESIS_SYSTEM_PROMPT = """\
你是一位严谨的通用学术开题报告写作助手。你将收到三类输入：
1. Gemini Deep Research 生成的课题底稿；
2. 本地参考文献库中提取的摘要、结论或正文片段；
3. 目标学校或单位的开题报告模板章节结构。

请严格执行：
- 先识别新课题的研究对象、核心科学问题、研究目标、技术路线与验证方案。
- 必须沿用模板中的章节标题和层级结构，不擅自改名、增删或重排核心章节。
- “国内外研究现状”必须采用夹叙夹议的学术批判性语调：先概括研究趋势，再结合文献指出方法、证据、局限和启示；禁止机械罗列“某某做了什么”。
- 将本地文献内容无缝缝合到 Gemini 底稿逻辑中，正文引用按出场顺序连续编号，参考文献列表与正文编号一一对应。
- 字数分配要均衡，背景、现状、评述、研究方案和实验验证都要有足够信息量。
- 技术路线部分要采用“问题引出 -> 实施路径 -> 方法/模型 -> 参数或指标解释 -> 本课题意义”的连贯叙事。
- 所有公式使用纯净 Markdown LaTeX：行内 $...$，独立公式 $$...$$。
- 输出完整 Markdown 正文，不要解释生成过程。
"""


PPT_BLUEPRINT_PROMPT = """\
请基于开题报告正文生成 15-16 页学术答辩 PPT 制作与演讲蓝图。
每页必须包含：
【幻灯片标题】
【指定插入图片】
【PPT 核心文字】
【汇报讲稿】

通用结构：
1. 封面：自动提取题目、汇报人、单位或导师信息；
2. 目录大纲；
3. 研究背景；
4. 核心痛点或科学问题；
5. 国内外现状与局限；
6. 研究目标与技术路线；
7-10. 研究内容模块，按正文自动拆分为 3-4 页；
11. 实验方案或验证机制；
12. 创新点；
13. 研究基础；
14. 预期成果；
15. 进度安排；
16. 结尾页。

图片分配规则：
- 根据图片文件名语义自动分配至最匹配页面；
- 若未找到合适图片，保留红色占位符；
- 每页讲稿约 100 字，口语化、自信、逻辑自然。
"""


SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".docx", ".doc", ".pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


@dataclass
class SourceBundle:
    draft_text: str
    template_structure: list[str]
    references_pack: str


def _path(value: str | os.PathLike) -> Path:
    return Path(value).expanduser().resolve()


def _safe_read_text(path: Path, limit: int | None = None) -> str:
    data = path.read_text(encoding="utf-8", errors="ignore")
    return data[:limit] if limit else data


def _read_docx_text(path: Path, limit: int | None = None) -> str:
    doc = Document(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                chunks.append(row_text)
    text = "\n".join(chunks)
    return text[:limit] if limit else text


def _read_csv_text(path: Path, limit: int | None = None) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for idx, row in enumerate(reader):
            if idx > 120:
                break
            clean = [cell.strip() for cell in row if cell.strip()]
            if clean:
                rows.append(" | ".join(clean))
    text = "\n".join(rows)
    return text[:limit] if limit else text


def _read_pdf_text(path: Path, limit: int | None = None) -> str:
    try:
        import pypdf  # type: ignore
    except Exception:
        return f"[PDF 文献占位：{path.name}。当前环境未安装 pypdf，未提取正文。]"
    chunks: list[str] = []
    with path.open("rb") as f:
        reader = pypdf.PdfReader(f)
        for page in reader.pages[:8]:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text.strip())
    text = "\n".join(chunks)
    return text[:limit] if limit else text


def _read_any_text(path: Path, limit: int | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_text(path, limit)
    if suffix == ".csv":
        return _read_csv_text(path, limit)
    if suffix == ".pdf":
        return _read_pdf_text(path, limit)
    return _safe_read_text(path, limit)


def _iter_reference_paths(references_dir: str | os.PathLike) -> list[Path]:
    root = _path(references_dir)
    if root.is_file():
        if root.suffix.lower() in SUPPORTED_TEXT_SUFFIXES:
            return [root]
        return []
    if not root.exists():
        raise FileNotFoundError(f"Reference path not found: {root}")
    paths = [
        p
        for p in root.rglob("*")
        if p.is_file() and not p.name.startswith("~$") and p.suffix.lower() in SUPPORTED_TEXT_SUFFIXES
    ]
    return sorted(paths, key=lambda p: p.name.lower())


def _build_references_pack(references_dir: str | os.PathLike, max_files: int = 80, chars_per_file: int = 2200) -> str:
    entries: list[str] = []
    for idx, path in enumerate(_iter_reference_paths(references_dir)[:max_files], start=1):
        text = _read_any_text(path, chars_per_file).strip()
        if not text:
            continue
        entries.append(f"[R{idx}] 文件名：{path.name}\n{text}")
    return "\n\n".join(entries) if entries else "未从 references_dir 中提取到可读文献内容。"


def _extract_template_structure(template_file: str | os.PathLike) -> list[str]:
    path = _path(template_file)
    text = _read_any_text(path, 20000)
    candidates: list[str] = []
    heading_patterns = [
        r"^第[一二三四五六七八九十]+[部分章节].{0,40}$",
        r"^[一二三四五六七八九十]+、.{1,60}$",
        r"^\d+(?:\.\d+){0,3}\s+.{1,80}$",
        r"^主要参考文献$",
        r"^参考文献$",
        r"^经费预算$",
        r"^研究基础$",
    ]
    for raw in text.splitlines():
        line = re.sub(r"\s+", " ", raw.strip())
        if not line or len(line) > 100:
            continue
        if any(re.match(pattern, line) for pattern in heading_patterns) and line not in candidates:
            candidates.append(line)
    return candidates or DEFAULT_TEMPLATE_STRUCTURE.copy()


def _call_openai(prompt: str, model: str | None = None, timeout: int = 900) -> str:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set.")
    payload = {
        "model": model or os.getenv("KAITI_OPENAI_MODEL", "gpt-4.1"),
        "input": prompt,
    }
    req = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        data = json.loads(resp.read().decode("utf-8"))
    chunks: list[str] = []
    for item in data.get("output", []):
        for content in item.get("content", []):
            if content.get("type") in {"output_text", "text"}:
                chunks.append(content.get("text", ""))
    result = "\n".join(chunks).strip()
    if not result:
        raise RuntimeError("OpenAI response did not contain output text.")
    return result


def _build_source_bundle(
    gemini_draft_file: str | os.PathLike,
    references_dir: str | os.PathLike,
    template_file: str | os.PathLike,
) -> SourceBundle:
    draft = _read_any_text(_path(gemini_draft_file), 70000)
    refs = _build_references_pack(references_dir)
    structure = _extract_template_structure(template_file)
    return SourceBundle(draft_text=draft, template_structure=structure, references_pack=refs)


def synthesize_and_rewrite_proposal(
    gemini_draft_file: str | os.PathLike,
    references_dir: str | os.PathLike,
    template_file: str | os.PathLike,
    output_file: str | os.PathLike,
) -> Path:
    """Create a template-aligned proposal from a Gemini draft and local references.

    The function attempts to call an LLM through the OpenAI Responses API. If no
    API key is configured, it writes a complete, reusable prompt package to
    `output_file` rather than silently inventing content.
    """

    bundle = _build_source_bundle(gemini_draft_file, references_dir, template_file)
    prompt = (
        f"{SYNTHESIS_SYSTEM_PROMPT}\n\n"
        "# 模板章节结构\n"
        + "\n".join(f"- {item}" for item in bundle.template_structure)
        + "\n\n# Gemini Deep Research 底稿\n"
        + bundle.draft_text
        + "\n\n# 本地文献库提取内容\n"
        + bundle.references_pack
    )
    try:
        output = _call_openai(prompt)
    except Exception as exc:
        output = (
            "# 通用学术开题报告生成任务包\n\n"
            f"> 当前未完成在线 LLM 生成：{exc}\n\n"
            "## 可复用生成指令\n\n"
            f"{prompt}\n"
        )
    dst = _path(output_file)
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(output, encoding="utf-8")
    return dst


def _set_rfonts(element, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key, value in {
        "w:ascii": latin,
        "w:hAnsi": latin,
        "w:cs": latin,
        "w:eastAsia": east_asia,
    }.items():
        rfonts.set(qn(key), value)


def _format_run(run, east_asia: str, latin: str, size_pt: float | None, bold: bool | None) -> None:
    run.font.name = latin
    _set_rfonts(run._element, east_asia, latin)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _apply_page_border(section, template_border=None) -> None:
    sect_pr = section._sectPr
    for old in list(sect_pr.findall(qn("w:pgBorders"))):
        sect_pr.remove(old)
    if template_border is not None:
        border = deepcopy(template_border)
    else:
        border = OxmlElement("w:pgBorders")
        border.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            child = OxmlElement(f"w:{edge}")
            child.set(qn("w:val"), "single")
            child.set(qn("w:sz"), "8")
            child.set(qn("w:space"), "24")
            child.set(qn("w:color"), "000000")
            border.append(child)
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is not None:
        sect_pr.insert(sect_pr.index(pg_mar), border)
    else:
        sect_pr.append(border)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        elem = OxmlElement(f"w:{edge}")
        if edge in {"top", "bottom"}:
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "12")
            elem.set(qn("w:color"), "000000")
        else:
            elem.set(qn("w:val"), "nil")
            elem.set(qn("w:sz"), "0")
            elem.set(qn("w:color"), "auto")
        borders.append(elem)


def _set_cell_borders(cell, header_bottom: bool = False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        old = borders.find(qn(f"w:{edge}"))
        if old is not None:
            borders.remove(old)
        elem = OxmlElement(f"w:{edge}")
        if edge == "bottom" and header_bottom:
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "6")
            elem.set(qn("w:color"), "000000")
        else:
            elem.set(qn("w:val"), "nil")
            elem.set(qn("w:sz"), "0")
            elem.set(qn("w:color"), "auto")
        borders.append(elem)


def _is_caption(text: str) -> bool:
    return bool(re.match(r"^(图|表)\s*[\d一二三四五六七八九十]+", text.strip()))


def format_academic_docx(
    input_docx: str | os.PathLike,
    template_file: str | os.PathLike,
    output_docx: str | os.PathLike,
) -> Path:
    """Apply cross-disciplinary Chinese academic DOCX formatting."""

    src = _path(input_docx)
    template_path = _path(template_file)
    dst = _path(output_docx)
    doc = Document(str(src))
    template = Document(str(template_path))
    tsec = template.sections[0]
    template_border = tsec._sectPr.find(qn("w:pgBorders"))
    template_border = deepcopy(template_border) if template_border is not None else None

    for section in doc.sections:
        section.page_width = tsec.page_width
        section.page_height = tsec.page_height
        section.top_margin = tsec.top_margin
        section.bottom_margin = tsec.bottom_margin
        section.left_margin = tsec.left_margin
        section.right_margin = tsec.right_margin
        _apply_page_border(section, template_border)

    in_refs = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in {"主要参考文献", "参考文献"}:
            in_refs = True
        elif re.match(r"^(第二部分|二、|2\s+|2\.)", text):
            in_refs = False

        if para.style.name.startswith("Heading 1") or re.match(r"^(第[一二三四五六七八九十]+[部分章节]|[一二三四五六七八九十]+、)", text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                _format_run(run, "黑体", "Times New Roman", 15, True)
        elif para.style.name.startswith("Heading 2") or re.match(r"^\d+\.\d+\s+", text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                _format_run(run, "黑体", "Times New Roman", 14, True)
        elif para.style.name.startswith("Heading 3") or re.match(r"^\d+\.\d+\.\d+\s+", text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                _format_run(run, "黑体", "Times New Roman", 12, True)
        elif _is_caption(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                _format_run(run, "宋体", "Times New Roman", 10.5, True)
        elif in_refs and re.match(r"^\[\d+\]", text):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.first_line_indent = -Cm(0.74)
            para.paragraph_format.line_spacing = 1.25
            for run in para.runs:
                _format_run(run, "宋体", "Times New Roman", 10.5, False)
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                _format_run(run, "宋体", "Times New Roman", 12, False)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _set_table_borders(table)
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                _set_cell_borders(cell, header_bottom=row_idx == 0)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = None
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        _format_run(run, "宋体", "Times New Roman", 10.5, row_idx == 0)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return dst


def _extract_title_and_people(text: str) -> dict[str, str]:
    compact = re.sub(r"\s+", " ", text)
    title = ""
    for pattern in [
        r"论文题目[:：]?\s*([^。；;\n]{8,120})",
        r"题目[:：]?\s*([^。；;\n]{8,120})",
        r"研究题目[:：]?\s*([^。；;\n]{8,120})",
    ]:
        match = re.search(pattern, compact)
        if match:
            title = match.group(1).strip()
            break
    if not title:
        for line in text.splitlines():
            if 12 <= len(line.strip()) <= 80 and not re.match(r"^\d", line.strip()):
                title = line.strip()
                break
    title = title or "学术开题报告"
    presenter = re.search(r"(姓名|汇报人)[:：]?\s*([\u4e00-\u9fa5A-Za-z0-9·]{2,20})", compact)
    advisor = re.search(r"(导师|指导教师)[:：]?\s*([^。；;\n]{2,30})", compact)
    unit = re.search(r"(学院|单位|学校)[:：]?\s*([^。；;\n]{2,40})", compact)
    return {
        "title": title,
        "presenter": presenter.group(2).strip() if presenter else "请补充",
        "advisor": advisor.group(2).strip() if advisor else "请补充",
        "unit": unit.group(2).strip() if unit else "请补充",
    }


def _extract_section_text(text: str, keywords: Sequence[str], limit: int = 260) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n+", text) if p.strip()]
    hits = [p for p in paragraphs if any(k in p for k in keywords)]
    if not hits:
        hits = paragraphs[:3]
    merged = "；".join(hits[:3])
    return merged[:limit]


def _extract_research_modules(text: str) -> list[str]:
    modules: list[str] = []
    for line in text.splitlines():
        clean = re.sub(r"\s+", " ", line.strip())
        if not clean:
            continue
        if re.match(r"^(\d+\.\d+(\.\d+)?|[（(]?\d+[）)]|[一二三四]、)", clean) and any(
            key in clean for key in ("研究内容", "内容", "方法", "模型", "实验", "方案", "技术", "系统")
        ):
            if clean not in modules and len(clean) <= 90:
                modules.append(clean)
    if len(modules) < 3:
        modules += ["研究内容一：理论基础与问题建模", "研究内容二：方法体系与模型构建", "研究内容三：实验验证与应用评价"]
    return modules[:4]


def _collect_images(image_list: str | os.PathLike | Iterable[str | os.PathLike]) -> list[Path]:
    if isinstance(image_list, (str, os.PathLike)):
        raw = _path(image_list)
        if raw.is_dir():
            return sorted([p for p in raw.iterdir() if p.suffix.lower() in IMAGE_SUFFIXES], key=lambda p: p.name.lower())
        if raw.is_file() and raw.suffix.lower() in IMAGE_SUFFIXES:
            return [raw]
        if raw.is_file():
            return [
                _path(line.strip())
                for line in raw.read_text(encoding="utf-8", errors="ignore").splitlines()
                if line.strip()
            ]
        parts = [part.strip() for part in str(image_list).split(";") if part.strip()]
        return [_path(part) for part in parts]
    return [_path(item) for item in image_list]


def _score_image(name: str, keywords: Sequence[str]) -> int:
    lower = name.lower()
    return sum(3 if key.lower() in lower else 0 for key in keywords)


def _assign_images(images: list[Path]) -> dict[int, Path]:
    slide_keywords = {
        4: ["痛点", "问题", "背景", "矛盾", "需求", "pain", "problem"],
        6: ["路线", "技术路线", "框架", "roadmap", "route", "workflow"],
        7: ["模型", "方法", "仿真", "机制", "model", "method"],
        8: ["检测", "感知", "数据", "指标", "评价", "sensor", "data"],
        9: ["算法", "网络", "降阶", "预测", "ai", "ml", "algorithm"],
        10: ["控制", "闭环", "优化", "决策", "control", "optimization"],
        11: ["实验", "平台", "装置", "验证", "experiment", "platform"],
        13: ["基础", "成果", "预研", "结果", "foundation", "result"],
        15: ["进度", "计划", "时间", "甘特", "schedule", "gantt"],
    }
    available = images[:]
    assigned: dict[int, Path] = {}
    for slide, keys in slide_keywords.items():
        if not available:
            break
        best = max(available, key=lambda p: _score_image(p.name, keys))
        if _score_image(best.name, keys) > 0:
            assigned[slide] = best
            available.remove(best)
    return assigned


def _make_core_points(source: str, fallback: Sequence[str]) -> list[str]:
    text = re.sub(r"\s+", " ", source).strip()
    clauses = re.split(r"[。；;]", text)
    points = [c.strip() for c in clauses if 8 <= len(c.strip()) <= 42]
    return (points[:3] or list(fallback))[:4]


def _build_generic_slides(text: str, images: list[Path]) -> list[dict]:
    meta = _extract_title_and_people(text)
    modules = _extract_research_modules(text)
    assignment = _assign_images(images)
    background = _extract_section_text(text, ["背景", "意义", "需求", "发展"])
    pain = _extract_section_text(text, ["问题", "痛点", "不足", "挑战", "矛盾"])
    target = _extract_section_text(text, ["目标", "技术路线", "路线", "任务"])
    experiment = _extract_section_text(text, ["实验", "验证", "评价", "样品", "平台"])
    innovation = _extract_section_text(text, ["创新", "特色", "贡献"])
    foundation = _extract_section_text(text, ["研究基础", "已有", "前期", "条件"])
    schedule = _extract_section_text(text, ["进度", "计划", "安排", "时间"])

    slides = [
        {
            "no": 1,
            "title": f"封面｜{meta['title']}",
            "core": [meta["unit"], "学术开题答辩", f"汇报人：{meta['presenter']}", f"导师：{meta['advisor']}"],
            "script": f"各位老师好，我汇报的题目是《{meta['title']}》。接下来我将围绕研究背景、目标路线、研究内容和预期安排进行说明。",
        },
        {
            "no": 2,
            "title": "目录大纲｜Contents",
            "core": ["01 研究背景与问题提出", "02 研究目标与技术路线", "03 研究内容与实验方案", "04 研究基础与进度安排"],
            "script": "本次汇报按照四个部分展开：先说明课题为什么值得做，再介绍研究目标和总体路线，随后展开关键研究内容，最后汇报基础条件和进度安排。",
        },
        {
            "no": 3,
            "title": "1.1 研究背景",
            "core": _make_core_points(background, ["研究对象具有明确应用需求", "现有方法存在提升空间", "课题具有理论与实践价值"]),
            "script": "本课题来源于当前领域发展的真实需求。现有研究和工程实践已经积累了一定基础，但面对更高精度、更强可靠性和更好应用效果的要求，仍需要进一步系统化研究。",
        },
        {
            "no": 4,
            "title": "1.2 核心痛点与科学问题",
            "core": _make_core_points(pain, ["关键状态难以准确刻画", "多目标之间存在冲突", "现有方法缺乏闭环验证"]),
            "script": "本课题要解决的核心问题，是现有方法在真实场景中仍存在信息不完整、机制解释不足和验证链条不闭合等问题。因此，研究不能只停留在单一指标改进，而要形成可解释、可验证的整体方案。",
        },
        {
            "no": 5,
            "title": "1.3 国内外现状与局限",
            "core": ["已有研究提供理论基础", "方法体系仍较分散", "面向本课题对象的系统整合不足"],
            "script": "从国内外现状看，相关研究已经在理论方法、实验手段和应用验证方面形成基础。但这些成果往往分散在不同方向，针对本课题对象的系统整合、关键参数解释和工程化验证仍显不足。",
        },
        {
            "no": 6,
            "title": "2.1 研究目标与总体路线",
            "core": _make_core_points(target, ["构建完整技术路线", "形成可验证的方法体系", "实现理论模型与应用场景衔接"]),
            "script": "基于上述问题，本课题拟建立一套从理论分析、方法构建到实验验证的完整路线。核心目标不是孤立完成某个环节，而是让模型、数据、实验和评价之间形成闭环。",
        },
    ]
    for idx, module in enumerate(modules[:4], start=7):
        slides.append(
            {
                "no": idx,
                "title": f"3.{idx - 6} {module}",
                "core": ["明确研究对象与输入输出", "构建关键方法或模型", "形成可评价的结果指标"],
                "script": f"这一部分围绕“{module}”展开。重点是把研究对象、关键变量和评价指标说清楚，并通过合适的方法或模型建立可分析、可复现的研究路径。",
            }
        )
    while len(slides) < 10:
        no = len(slides) + 1
        slides.append(
            {
                "no": no,
                "title": f"3.{no - 6} 研究内容补充模块",
                "core": ["补充关键技术环节", "完善验证与评价逻辑", "支撑总体研究目标"],
                "script": "本页用于补充说明研究中的关键技术环节，保证整体方案从问题提出到结果评价之间逻辑完整、层次清楚。",
            }
        )
    slides += [
        {
            "no": 11,
            "title": "3.5 实验方案与验证机制",
            "core": _make_core_points(experiment, ["设计对照实验", "建立评价指标", "形成数据-模型-结果验证闭环"]),
            "script": "实验方案将围绕研究目标设置对照和验证环节。通过关键指标采集、统计分析和模型结果对比，判断所提出方法是否真正改善了问题。",
        },
        {
            "no": 12,
            "title": "3.6 创新点与特色",
            "core": _make_core_points(innovation, ["理论机制创新", "方法体系创新", "验证路径创新"]),
            "script": "本课题的创新性主要体现在三方面：一是围绕具体问题建立解释机制；二是形成可执行的方法体系；三是通过实验或案例验证，使研究结论具有可落地性。",
        },
        {
            "no": 13,
            "title": "4.1 研究基础",
            "core": _make_core_points(foundation, ["具备相关研究积累", "具备数据或实验条件", "具备继续推进的技术基础"]),
            "script": "前期已经具备一定研究基础，包括资料积累、方法准备和实验条件。这些基础能够支撑课题后续从方案设计进入系统实施。",
        },
        {
            "no": 14,
            "title": "4.2 预期成果",
            "core": ["形成开题报告与技术路线", "形成模型或方法原型", "形成实验数据、论文或应用成果"],
            "script": "预期成果包括三类：第一是完整的理论与技术路线；第二是可复现的方法或系统原型；第三是支撑论文写作和后续应用的数据与结果。",
        },
        {
            "no": 15,
            "title": "4.3 进度安排",
            "core": _make_core_points(schedule, ["阶段一：资料整理与方案细化", "阶段二：方法构建与实验验证", "阶段三：结果分析与论文撰写"]),
            "script": "进度安排上，前期重点完成资料整理和方案细化，中期集中开展方法构建和实验验证，后期进行结果分析、论文撰写和成果凝练。",
        },
        {
            "no": 16,
            "title": "结尾页｜敬请批评指正",
            "core": ["谢谢各位老师", "敬请批评指正"],
            "script": "以上就是我的开题汇报。恳请各位老师对研究问题、技术路线和实验安排提出宝贵意见，我会根据建议继续完善课题设计。谢谢大家。",
        },
    ]
    for slide in slides:
        slide["image"] = assignment.get(slide["no"])
    return slides[:16]


def _write_blueprint_docx(slides: list[dict], output: Path) -> None:
    doc = Document()
    doc.add_heading("PPT制作与演讲蓝图指导书", 0)
    for slide in slides:
        doc.add_heading(f"Slide {slide['no']:02d}", level=1)
        doc.add_paragraph(f"【幻灯片标题】：{slide['title']}")
        image_path = slide.get("image")
        if image_path:
            doc.add_paragraph(f"【指定插入图片】：{image_path}")
        else:
            p = doc.add_paragraph()
            p.add_run("【指定插入图片】：")
            run = p.add_run("【红色占位符】请根据本页主题补充合适图片或图表")
            run.font.color.rgb = RGBColor(192, 0, 0)
        doc.add_paragraph("【PPT 核心文字】")
        for item in slide["core"]:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph(f"【汇报讲稿】：{slide['script']}")
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = 1.25
        for run in para.runs:
            _format_run(run, "宋体", "Times New Roman", 10.5, None)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def _write_blueprint_markdown(slides: list[dict], output: Path) -> None:
    lines = ["# PPT制作与演讲蓝图指导书", ""]
    for slide in slides:
        image_text = str(slide.get("image")) if slide.get("image") else "【红色占位符】请根据本页主题补充合适图片或图表"
        lines += [
            f"## Slide {slide['no']:02d}",
            "",
            f"**【幻灯片标题】**：{slide['title']}",
            "",
            f"**【指定插入图片】**：{image_text}",
            "",
            "**【PPT 核心文字】**",
            "",
        ]
        lines += [f"- {item}" for item in slide["core"]]
        lines += ["", f"**【汇报讲稿】**：{slide['script']}", ""]
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines), encoding="utf-8")


def generate_generic_ppt_blueprint(
    input_docx: str | os.PathLike,
    image_list: str | os.PathLike | Iterable[str | os.PathLike],
    output_blueprint: str | os.PathLike,
) -> Path:
    """Generate a generic 15-16 page academic defense PPT Word guide."""

    src = _path(input_docx)
    output = _path(output_blueprint)
    text = _read_docx_text(src, 70000)
    images = _collect_images(image_list)
    slides = _build_generic_slides(text, images)
    if output.suffix.lower() == ".docx":
        _write_blueprint_docx(slides, output)
    else:
        _write_blueprint_markdown(slides, output)
    return output


def rewrite_and_balance_proposal(
    gemini_draft_file: str | os.PathLike,
    output_file: str | os.PathLike,
    references_dir: str | os.PathLike | None = None,
    template_file: str | os.PathLike | None = None,
) -> Path:
    """Backward-compatible wrapper for older callers.

    New code should call `synthesize_and_rewrite_proposal(...)` directly.
    """

    if references_dir is None or template_file is None:
        draft = _path(gemini_draft_file)
        dst = _path(output_file)
        prompt = (
            "# 通用学术开题报告生成任务包\n\n"
            "请补充 references_dir 与 template_file 后调用 synthesize_and_rewrite_proposal。\n\n"
            "## 底稿内容\n\n"
            f"{_read_any_text(draft, 70000)}\n"
        )
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_text(prompt, encoding="utf-8")
        return dst
    return synthesize_and_rewrite_proposal(gemini_draft_file, references_dir, template_file, output_file)


def generate_ppt_blueprint(
    input_docx: str | os.PathLike,
    image_dir: str | os.PathLike,
    output_blueprint: str | os.PathLike,
) -> Path:
    """Backward-compatible wrapper for older callers."""

    return generate_generic_ppt_blueprint(input_docx, image_dir, output_blueprint)


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generic academic proposal workflow helpers.")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("synthesize")
    p.add_argument("gemini_draft_file")
    p.add_argument("references_dir")
    p.add_argument("template_file")
    p.add_argument("output_file")

    p = sub.add_parser("format")
    p.add_argument("input_docx")
    p.add_argument("template_file")
    p.add_argument("output_docx")

    p = sub.add_parser("blueprint")
    p.add_argument("input_docx")
    p.add_argument("image_list")
    p.add_argument("output_blueprint")

    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    if args.cmd == "synthesize":
        out = synthesize_and_rewrite_proposal(
            args.gemini_draft_file, args.references_dir, args.template_file, args.output_file
        )
    elif args.cmd == "format":
        out = format_academic_docx(args.input_docx, args.template_file, args.output_docx)
    elif args.cmd == "blueprint":
        out = generate_generic_ppt_blueprint(args.input_docx, args.image_list, args.output_blueprint)
    else:  # pragma: no cover
        raise SystemExit(f"Unknown command: {args.cmd}")
    print(out)


if __name__ == "__main__":
    main()
